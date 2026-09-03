"""Generate the presentation Word document — no numbers, natural speech."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

def para(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.7)
    return p

def say(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)

def cmd_line(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Consolas"
    run.font.color.rgb = RGBColor(80, 80, 80)

def output(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.size = Pt(9)
        run.font.name = "Consolas"
        run.font.color.rgb = RGBColor(60, 60, 60)

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("· · ·")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(180, 180, 180)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ═══════════════════════════════════════════════════════════════════
say("先跑第一个命令，pcap，深度抓包。")
cmd_line("python -m src.cli.main pcap cloudflare.com")

para("这个命令做的事情很简单：在浏览器打开 cloudflare 的同时，用 OpenSSL 发起一个 TLS 连接，把握手过程中收发的每一条消息以十六进制原始数据的形式抓下来。大家看输出，从上到下分成三个层次。")

output([
    "TLS Handshake Capture: cloudflare.com:443",
    "Method: openssl_msg",
    "Cipher: TLS_AES_256_GCM_SHA384",
    "Protocol: TLSv1.3",
])

para("最上面是概览。握手结束之后，程序从输出里提取了两个关键信息：协商出来的密码套件和 TLS 协议版本。密码套件是 TLS_AES_256_GCM_SHA384，这个字符串包含了四部分信息——密钥交换走的是 ECDHE，也就是椭圆曲线临时密钥交换；身份认证由服务器证书决定；对称加密用 AES，GCM 模式自带防篡改；哈希算法 SHA-384 用于派生会话密钥和做完整性校验。协议版本是 TLS 1.3，目前最新的标准。")

output([
    "   1. >>> ClientHello            TLS 1.3",
    "   2. <<< ServerHello            TLS 1.3",
    "   6. <<< Certificate            TLS 1.3",
    "   7. <<< CertificateVerify      TLS 1.3",
    "   8. <<< Finished               TLS 1.3",
])

para("中间是逐条记录。ClientHello 是客户端先开口，告诉服务器我能用这些算法、这些 TLS 版本，带一个随机数。ServerHello 是服务器选好算法、也带一个随机数回过来。第六条 Certificate 是这里面最大的一条，包含服务器的 X.509 证书链，DER 编码的二进制数据。第七条 CertificateVerify，服务器用自己的私钥签了一个名，证明它确实持有这个证书。第八条 Finished，双方各自对握手过程中所有消息算一个 HMAC，确认握手没有被中间人篡改过。可以看到，TLS 1.3 的握手是非常紧凑的——从 ClientHello 到 Finished 只走了两个往返。")

output([
    "Record 1: >>> Handshake (ClientHello)",
    "  Hex: 01 00 06 09 03 03 ab 78 2b 31 cb 67 7b 9d df b7 ...",
])

para("最下面是每条记录的 hex 原始数据。第一个字节 0x01 就是 HandshakeType——ClientHello 的编号。后面紧跟着长度字段、TLS 版本、随机数、会话 ID、密码套件列表、扩展列表……每一条消息都是严格按照 RFC 8446 规定的二进制结构编码的。这就是 TLS 协议的最底层——本质上就是按照固定的字节偏移去读数据。")

para("所以 pcap 让我们看到了 TLS 握手的完整骨架。但全是 hex，看不出这个网站是否抗量子。下一步就要从这些原始数据里提取关键字段——服务器在 ServerHello 里到底选了哪个密钥交换组？")

divider()

say("第二个命令，detect，PQC 握手检测。")
cmd_line("python -m src.cli.main detect cloudflare.com")

para("它的原理是：我们在发起 TLS 连接时，在 ClientHello 的 supported_groups 扩展中主动声明多个密钥交换组，把抗量子的排在第一位。具体来说，告诉服务器：我会 X25519MLKEM768——这是一个混合方案，在传统的 X25519 外面包了一层 ML-KEM-768 抗量子算法——还会 x25519 和 secp256r1 两个经典算法。注意这个顺序是有讲究的：TLS 1.3 规定服务器必须按客户端给出的优先级来选，所以把 PQC 组排在第一位，如果服务器支持，它就会选这个，不会跳过它去选后面的经典组。反过来，如果服务器选了经典组，只有一种可能——它不支持 PQC。然后程序去解析 ServerHello 的 key_share 扩展，看服务器到底选了哪个。")

para("解析的过程是逐字节扫描。ServerHello 消息体的开头是 HandshakeType、长度、版本、随机数、会话 ID、密码套件、压缩方法……这些都跳过。然后从扩展区开始，每个扩展是类型加长度加数据。当扫到 type 等于 0x0033 的时候，这就是 key_share 扩展。里面第一个字段是两字节的组 ID——服务器选的那个。")

para("如果这个组 ID 是 0x11EC，查表就是 X25519MLKEM768，那服务器在客户端明确提供了 PQC 选项的情况下主动选了抗量子算法，这就是无可辩驳的直接证据。如果服务器选了 0x001D 也就是 x25519，说明客户端给了 PQC 选项但服务器没选，确认不支持。")

output([
    "PQC Detection Result: cloudflare.com:443",
    "============================================================",
    "  Protocol:       TLSv1.3",
    "  Cipher Suite:   TLS_AES_256_GCM_SHA384",
    "  PQC Supported:  ✓ YES",
    "  PQC Algorithm:  X25519MLKEM768",
    "  PQC Group ID:   0x11EC",
])

para("结果出来了。服务器选了 0x11EC，也就是 X25519MLKEM768，PQC 确认。key_share 的大小比纯经典的 32 字节大了很多，因为混合方案里同时包含了 X25519 公钥和 ML-KEM-768 密文，这也是判断是否 PQC 的一个辅助信号——PQC 握手的数据量明显更大。")

para("作为对比，再看看百度。")

cmd_line("python -m src.cli.main detect baidu.com")

output([
    "  Cipher Suite:  ECDHE-RSA-AES128-GCM-SHA256",
    "  PQC Supported: ✗ NO",
])

para("百度的 ServerHello 里 key_share 选的是经典 x25519，PQC 不支持。而且注意它的密码套件名字格式不一样——ECDHE-RSA-AES128-GCM-SHA256，五个字段用横线连起来，这是 TLS 1.2 的命名规范，跟刚才 TLS 1.3 的格式完全不同。TLS 1.3 简化了命名，去掉了密钥交换和认证字段，因为这两个已经不由密码套件名字来规定了。")

para("所以 detect 回答的是：这个网站的传输层支不支持抗量子密钥交换。但它回答不了另一个问题——证书。证书的签名算法是什么？公钥是什么类型？这是 cert 命令要做的。")

divider()

say("第三个命令，cert，证书深度分析。")
cmd_line("python -m src.cli.main cert cloudflare.com")

para("这个命令拿到证书的方式很简单——Python 的 ssl 模块完成 TLS 握手后，调 getpeercert 方法，直接从服务器发来的 Certificate 消息里把 DER 编码的 X.509 证书取出来。然后交给 cryptography 库解析。不用发 HTTP 请求，握手完就拿证书，速度很快。")

output([
    "  使用者 (CN):    cloudflare.com",
    "  颁发者 (CN):    WE1",
    "  颁发机构:       Google Trust Services",
])

para("先看基本信息。这个证书是 Google Trust Services 签发给 cloudflare.com 的。注意有效期——现在的趋势是签发短期证书，减少密钥泄露的时间窗口，不像以前一签就是一年两年。")

output([
    "  签名算法:  ECDSA-SHA256",
    "  OID:       1.2.840.10045.4.3.2",
])

para("签名算法是 ECDSA-SHA256。OID 就是这个算法在国际标准里的身份证号——1.2.840.10045.4.3.2，全球唯一。我的程序内置了一张完整的 OID 映射表，覆盖了 RSA、ECDSA、SM2 国密、Dilithium、Falcon、SPHINCS+ 这些主流签名算法。判断逻辑是看 OID 前缀——以 1.2.156.10197 开头就是国密，以 1.3.6.1.4.1.2.267 或 1.3.9999 开头就是抗量子算法。当 CA 将来开始签发抗量子证书的时候，程序不用改任何代码就能自动识别，标注为 NIST 抗量子标准。ECDSA 是经典算法，不能抵抗量子攻击。")

output([
    "  公钥类型:  EC",
])

para("公钥这边，程序通过 Python 的 isinstance 判断类型——RSAPublicKey 就是 RSA，EllipticCurvePublicKey 就是椭圆曲线，Ed25519PublicKey 就是 Ed25519。如果出现这三种之外的类型，标记为 Unknown，很可能是将来部署的 PQC 公钥。Cloudflare 这里是椭圆曲线，程序还进一步提取了曲线的具体名称——比如 secp256r1 也就是 P-256 曲线。对每种曲线也都建立了 OID 映射表。")

output([
    "  抗量子安全:  ✗ 经典 (非抗量子)",
])

para("安全评估的判定逻辑很简单：签名算法是 PQC 的、或者公钥类型是 PQC 的、或者扩展域里有 PQC 相关的自定义 OID——三个条件满足任何一个，就判定为抗量子。Cloudflare 这三条都不满足，所以结论是经典，非抗量子。NIST 安全级别也一并给出，根据算法估算比特安全强度，但这个估算只对经典计算机有效——量子计算机跑 Shor 算法的时候，RSA 和 ECDSA 是多项式时间可破的，再高的经典安全级别也没用。")

output([
    "  证书扩展:",
    "    Key Usage, SAN, CRL Distribution Points ...",
    "    Custom OID 1.3.6.1.4.1.11129.2.4.2",
])

para("最后是扩展域。程序把它们分成两类：已知扩展直接查表显示名字——Key Usage 是公钥用途，SAN 是域名列表，CRL 是证书吊销列表地址。不在表里的自定义 OID，如果前缀落在 PQC 或国密的范围内，就单独计数——这可以作为辅助信号，万一将来有人在扩展域里嵌入 PQC 算法标识，这里能捕捉到。")

para("综合来看，Cloudflare 是一个很典型的案例：TLS 传输层已经用上了抗量子的密钥交换，但证书层还是经典算法。这其实反映了当前互联网的普遍现状——密钥交换可以先升级，因为它是纯技术选择，服务器端改个配置就行。但证书体系的升级涉及整个 PKI 基础设施：CA 要支持新算法、根证书要更新、浏览器要信任新的证书链……这个过程要慢得多。")

divider()

say("总结一下三个命令的关系。")

para("它们本质上都是从同一个 TLS 握手中抓数据，只是提取的层面不同。pcap 最底层，拿到全部 hex 记录，数据量最大但信息密度最低。detect 从 ClientHello 和 ServerHello 里提取 key_share 扩展，判断密钥交换是不是 PQC，回答传输层的问题。cert 从 Certificate 消息里取出 X.509 证书，解析签名算法和公钥类型，回答证书层的问题。三个命令合起来，从最原始的十六进制数据一路看到最上层的信任体系，对一个网站的密码安全状况做了一次完整的多层面体检。")

para("目前实测的结果是：Cloudflare 这类走在最前面的 CDN 已经在传输层做了混合 PQC，但证书层的 PQC 基本还是空白，大部分网站两个层面都没有 PQC。这个项目在设计上考虑到了扩展性——OID 映射表和 PQC 组 ID 映射表都已经预置了 NIST 标准化的全部算法，新算法出现时只需要在表里加一行，不需要改逻辑。所以这个工具可以持续使用，来监测整个互联网从经典密码体系向抗量子密码体系迁移的过程。")

divider()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run("谢谢老师")
run.bold = True
run.font.size = Pt(14)

# ── Save ──
out = "C:/Users/26554/Desktop/pqc-https-project/汇报稿-演示-v2.docx"
doc.save(out)
print(f"Saved: {out}")
